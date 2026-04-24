#!/usr/bin/env python3
"""
BCBS VT Proposal Report Reference-Doc Builder

Builds ``data/letterhead/assets/reference-proposal-report.docx`` from the
authoritative Blue Cross VT Proposal Report Portrait Template
(``proposal-report-template.dotx``). This is the pandoc ``--reference-doc``
used by ``build-letterhead.sh`` for memos, reports, proposals, strategy
documents, and polished BCBS Word deliverables.

Design principle: meet pandoc where it writes.

Pandoc emits paragraph style IDs drawn from its own vocabulary, not from
Word's defaults. If the reference doc does not define those style IDs,
Word silently falls back to Normal and the output looks unstyled. The
pandoc styleIds we must cover are:

  Title, Subtitle, Author, Date            (YAML frontmatter block)
  Heading1, Heading2, Heading3, Heading4   (markdown # / ## / ### / ####)
  BodyText, FirstParagraph, Compact        (body + tight list + table cells)
  Caption, IntenseQuote                    (image captions, blockquotes)
  Table (table-type style)                 (markdown tables)

The ``Heading1`` style is load-bearing: the official .dotx defines it as
bold ALL-CAPS navy (#00355E) on a light-blue (#99D6EA) shading band, and
that is the visual signature of a BCBS document. This script DOES NOT
TOUCH Heading1 — the template wins.

For every other style we either:
  - upsert a paragraph style whose size, color, and font match the two
    reference documents in ~/Documents/BCBS (Jira One-Sheet, Digital
    Infrastructure Strategy Memo V2), OR
  - inject a table-type style via raw XML, because python-docx's
    paragraph-style API can't describe table borders / firstRow
    conditional formatting.

The script is otherwise non-destructive: it preserves the template's
margins, Normal style, Heading1 style, Header, Footer, and ListParagraph
style, and it preserves the header (logo + blue header band).

Usage:   python3 style-proposal-report.py
Requires: pip install python-docx
"""

import io
import sys
import zipfile
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from lxml import etree
except ImportError:
    print("Error: python-docx / lxml not installed.")
    print("Run: pip install python-docx lxml")
    sys.exit(1)


# --- BCBS VT palette (authoritative Brand Style Guide 2025-10 + template) ---
# #00355E is the legacy dark navy carried by the official .dotx template's
# Heading1 band. The reference memos use the same hex for the Title. We
# match it here so Title and Heading1 look from the same family.
PRIMARY_BLUE = RGBColor(0x00, 0x77, 0xC8)   # Blue Cross VT Blue (PMS 3005 C)
DARK_NAVY    = RGBColor(0x00, 0x35, 0x5E)   # Template Heading1 + ref Title
MID_GRAY     = RGBColor(0x59, 0x59, 0x59)   # Byline / subtitle gray
CAPTION_GRAY = RGBColor(0x63, 0x66, 0x6A)   # Cool Gray 10 C
FOOTER_GRAY  = RGBColor(0x66, 0x66, 0x66)
BLACK        = RGBColor(0x1A, 0x1A, 0x1A)
BORDER_GRAY  = "BFBFBF"                      # Table cell borders (hex no '#')
HEADER_SHADE = "E6F2FA"                      # Light table header tint

# --- Typography ---
# The authoritative Writing and Tone Style Guide names Calibri / DIN 2014 /
# Arial as approved fonts; the .dotx itself is Calibri-based. The reference
# one-sheet and memo keep Calibri throughout and rely on size + color (not
# font family) for hierarchy. We follow that convention.
BODY_FONT = "Calibri"

SCRIPT_DIR  = Path(__file__).parent
ASSETS_DIR  = SCRIPT_DIR.parent / "assets"
SOURCE_DOTX = ASSETS_DIR / "proposal-report-template.dotx"
OUTPUT_PATH = ASSETS_DIR / "reference-proposal-report.docx"


# ---------------------------------------------------------------------------
# .dotx → .docx package conversion
# ---------------------------------------------------------------------------

def dotx_to_docx_bytes(dotx_path):
    """Rewrite the .dotx package as a .docx package (bytes).

    Swaps the single content-type override for ``word/document.xml`` from
    ``...template.main+xml`` to ``...document.main+xml``. All other parts
    (header, footer, media, styles, theme, settings) are copied verbatim.
    """
    template_ct = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.template.main+xml"
    )
    document_ct = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document.main+xml"
    )
    out_buf = io.BytesIO()
    with zipfile.ZipFile(dotx_path, "r") as zin:
        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        template_ct.encode("utf-8"),
                        document_ct.encode("utf-8"),
                    )
                zout.writestr(item, data)
    return out_buf.getvalue()


def strip_placeholder_body(doc):
    """Remove .dotx placeholder paragraphs so pandoc fills in real content.

    Preserves ``w:sectPr`` (section properties) so the header reference,
    footer reference, page size, and margins survive.
    """
    body = doc.element.body
    for child in list(body):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag in ("p", "tbl"):
            body.remove(child)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Paragraph-style upsert
# ---------------------------------------------------------------------------

def _find_style_by_id(doc, style_id):
    """Return the <w:style> element with given w:styleId, or None."""
    for s in doc.styles.element.findall(qn("w:style")):
        if s.get(qn("w:styleId")) == style_id:
            return s
    return None


def upsert_paragraph_style(
    doc,
    style_id,
    *,
    name=None,
    based_on="Normal",
    font=None,
    size_pt=None,
    bold=None,
    italic=None,
    caps=None,
    color=None,
    shading=None,
    space_before=None,
    space_after=None,
    line_spacing=None,
    align=None,
):
    """Create or update a paragraph style by styleId.

    pandoc writes w:pStyle by styleId, so we key on that rather than on
    the display name (python-docx's default indexer).

    All ``None`` fields are left at whatever the template / basedOn
    chain provides. That lets us stay non-destructive: a body paragraph
    gets Calibri 10pt 1.15 spacing from Normal without us restating it.
    """
    existing = _find_style_by_id(doc, style_id)
    if existing is not None:
        # Keep the style; clear the rPr/pPr we plan to overwrite, then
        # rewrite from the kwargs we were passed.
        for tag in ("w:pPr", "w:rPr"):
            for node in existing.findall(qn(tag)):
                existing.remove(node)
        style_el = existing
        action = "updated"
    else:
        style_el = OxmlElement("w:style")
        style_el.set(qn("w:type"), "paragraph")
        style_el.set(qn("w:styleId"), style_id)
        doc.styles.element.append(style_el)
        action = "added"

    # <w:name>
    if name is not None:
        nm = OxmlElement("w:name")
        nm.set(qn("w:val"), name)
        style_el.append(nm)
    # <w:basedOn>
    if based_on is not None:
        bo = OxmlElement("w:basedOn")
        bo.set(qn("w:val"), based_on)
        style_el.append(bo)
    # qFormat so Word surfaces it in the gallery
    qf = OxmlElement("w:qFormat")
    style_el.append(qf)

    # <w:pPr>
    ppr = OxmlElement("w:pPr")
    if shading is not None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), shading)
        ppr.append(shd)
    if space_before is not None or space_after is not None or line_spacing is not None:
        sp = OxmlElement("w:spacing")
        if space_before is not None:
            sp.set(qn("w:before"), str(int(Pt(space_before).twips)))
        if space_after is not None:
            sp.set(qn("w:after"), str(int(Pt(space_after).twips)))
        if line_spacing is not None:
            sp.set(qn("w:line"), str(int(line_spacing * 240)))
            sp.set(qn("w:lineRule"), "auto")
        ppr.append(sp)
    if align is not None:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), align)
        ppr.append(jc)
    if len(ppr):
        style_el.append(ppr)

    # <w:rPr>
    rpr = OxmlElement("w:rPr")
    if font is not None:
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), font)
        rf.set(qn("w:hAnsi"), font)
        rf.set(qn("w:cs"), font)
        rpr.append(rf)
    if bold is True:
        rpr.append(OxmlElement("w:b"))
    if italic is True:
        rpr.append(OxmlElement("w:i"))
    if caps is True:
        rpr.append(OxmlElement("w:caps"))
    if color is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "%02X%02X%02X" % (color[0], color[1], color[2]))
        rpr.append(c)
    if size_pt is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))
        rpr.append(sz)
        szcs = OxmlElement("w:szCs")
        szcs.set(qn("w:val"), str(int(size_pt * 2)))
        rpr.append(szcs)
    if len(rpr):
        style_el.append(rpr)

    print(f"  {action:7s}: {style_id}")


# ---------------------------------------------------------------------------
# Table-type style (BCBS grid)
# ---------------------------------------------------------------------------

TABLE_STYLE_XML = """
<w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         w:type="table" w:styleId="Table">
  <w:name w:val="Table"/>
  <w:basedOn w:val="TableNormal"/>
  <w:uiPriority w:val="59"/>
  <w:qFormat/>
  <w:pPr>
    <w:spacing w:before="40" w:after="40" w:line="240" w:lineRule="auto"/>
  </w:pPr>
  <w:rPr>
    <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:cs="Calibri"/>
    <w:sz w:val="20"/>
    <w:szCs w:val="20"/>
  </w:rPr>
  <w:tblPr>
    <w:tblBorders>
      <w:top    w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>
      <w:left   w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>
      <w:right  w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>
      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>
      <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>
    </w:tblBorders>
    <w:tblCellMar>
      <w:top w:w="80" w:type="dxa"/>
      <w:left w:w="120" w:type="dxa"/>
      <w:bottom w:w="80" w:type="dxa"/>
      <w:right w:w="120" w:type="dxa"/>
    </w:tblCellMar>
  </w:tblPr>
  <w:tblStylePr w:type="firstRow">
    <w:rPr>
      <w:b/>
      <w:color w:val="00355E"/>
    </w:rPr>
    <w:tcPr>
      <w:shd w:val="clear" w:color="auto" w:fill="E6F2FA"/>
    </w:tcPr>
  </w:tblStylePr>
</w:style>
""".strip()


def inject_table_style(doc):
    """Add the 'Table' table-type style pandoc references in <w:tblStyle>.

    python-docx's table-style API can't express tblBorders or the
    firstRow conditional formatting we want, so we inject the raw XML.
    Non-destructive: if the style already exists, replace its body.
    """
    existing = _find_style_by_id(doc, "Table")
    # Parse our string into the styles tree. etree.fromstring keeps w: ns.
    new_el = etree.fromstring(TABLE_STYLE_XML)
    if existing is not None:
        doc.styles.element.replace(existing, new_el)
        print("  updated: Table (table-type)")
    else:
        doc.styles.element.append(new_el)
        print("  added  : Table (table-type)")


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

def set_footer_if_empty(doc, text):
    """Populate every empty section footer with the required disclosure."""
    for section in doc.sections:
        footer = section.footer
        existing = "".join(p.text for p in footer.paragraphs).strip()
        if existing:
            print("  footer already populated; leaving untouched")
            continue
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = FOOTER_GRAY
        print("  footer added (centered, 8pt Calibri)")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("BCBS VT Proposal Report Reference-Doc Builder (v2)")
    print("=" * 52)
    print("Preserves: Normal, Heading1, Header, Footer, ListParagraph,")
    print("           margins, logo/banner header, the official template")
    print("           styles that already render correctly.")
    print("Adds/updates pandoc-expected styles: Title, Subtitle, Author,")
    print("           Date, Heading2-4, BodyText, FirstParagraph, Compact,")
    print("           Caption, IntenseQuote, and a Table table-style.")

    if not SOURCE_DOTX.exists():
        print(f"\nError: source template missing: {SOURCE_DOTX}")
        sys.exit(1)

    print(f"\nSource template: {SOURCE_DOTX.name}")
    docx_bytes = dotx_to_docx_bytes(SOURCE_DOTX)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_PATH.write_bytes(docx_bytes)
    print("  package converted: .dotx → .docx")

    doc = Document(str(OUTPUT_PATH))

    print("\nStripping placeholder body...")
    strip_placeholder_body(doc)

    print("\nUpserting paragraph styles pandoc writes...")

    # Title — direct-format pattern in the BCBS references: Calibri 15pt
    # bold dark navy. Pandoc emits <w:pStyle w:val="Title"/> when the
    # Markdown file has a YAML title.
    upsert_paragraph_style(
        doc, "Title", name="Title",
        font=BODY_FONT, size_pt=15, bold=True, color=DARK_NAVY,
        space_before=0, space_after=0, line_spacing=1.15,
    )
    # Subtitle — byline line, 8pt mid-gray, from the YAML subtitle field.
    upsert_paragraph_style(
        doc, "Subtitle", name="Subtitle",
        font=BODY_FONT, size_pt=8, bold=False, color=MID_GRAY,
        space_before=0, space_after=12, line_spacing=1.2,
    )
    # Author / Date — optional metadata paragraphs.
    upsert_paragraph_style(
        doc, "Author", name="Author",
        font=BODY_FONT, size_pt=10, color=CAPTION_GRAY, space_after=2,
    )
    upsert_paragraph_style(
        doc, "Date", name="Date",
        font=BODY_FONT, size_pt=10, color=CAPTION_GRAY, space_after=12,
    )
    # Heading2-4 — markdown ##/###/####. BCBS refs rarely use past H1;
    # we still need these defined so the look is coherent when someone
    # does nest. Calibri bold, BCBS primary blue, NO blue band (that's
    # reserved for H1).
    upsert_paragraph_style(
        doc, "Heading2", name="heading 2", based_on="Normal",
        font=BODY_FONT, size_pt=12, bold=True, color=DARK_NAVY,
        space_before=12, space_after=4, line_spacing=1.15,
    )
    upsert_paragraph_style(
        doc, "Heading3", name="heading 3", based_on="Normal",
        font=BODY_FONT, size_pt=11, bold=True, color=PRIMARY_BLUE,
        space_before=8, space_after=2, line_spacing=1.15,
    )
    upsert_paragraph_style(
        doc, "Heading4", name="heading 4", based_on="Normal",
        font=BODY_FONT, size_pt=11, bold=True, italic=True,
        color=PRIMARY_BLUE, space_before=6, space_after=2, line_spacing=1.15,
    )
    # BodyText / FirstParagraph — let them inherit Normal (Calibri 10pt,
    # 1.15 spacing, 5pt before / 10pt after) per the template. We only
    # define the style IDs so pandoc's pStyle references resolve; no
    # attribute overrides, so they visually match Normal.
    upsert_paragraph_style(
        doc, "BodyText", name="Body Text", based_on="Normal",
    )
    upsert_paragraph_style(
        doc, "FirstParagraph", name="First Paragraph", based_on="Normal",
        space_before=0,
    )
    # Compact — used in tight lists and inside table cells. Less vertical
    # space. Same font as Normal.
    upsert_paragraph_style(
        doc, "Compact", name="Compact", based_on="Normal",
        space_before=0, space_after=0, line_spacing=1.15,
    )
    # Caption / IntenseQuote — minor styles pandoc references.
    upsert_paragraph_style(
        doc, "Caption", name="caption", based_on="Normal",
        font=BODY_FONT, size_pt=9, italic=True, color=CAPTION_GRAY,
    )
    upsert_paragraph_style(
        doc, "IntenseQuote", name="Intense Quote", based_on="Normal",
        font=BODY_FONT, size_pt=10, italic=True, color=DARK_NAVY,
        space_before=8, space_after=8,
    )

    print("\nInjecting table-type style...")
    inject_table_style(doc)

    print("\nMargins: untouched (template sets 1\" all sides).")
    print("Heading1: untouched (template defines the blue-band look).")

    print("\nPopulating footer with Independent Licensee disclosure...")
    footer_text = (
        "Blue Cross and Blue Shield of Vermont is an independent licensee "
        "of the Blue Cross and Blue Shield Association.  |  "
        "bluecrossvt.org  |  Berlin, Vermont"
    )
    set_footer_if_empty(doc, footer_text)

    doc.save(str(OUTPUT_PATH))
    print(f"\nSaved: {OUTPUT_PATH}")
    print("\nUse with: build-letterhead.sh input.md [output.docx]")


if __name__ == "__main__":
    main()
