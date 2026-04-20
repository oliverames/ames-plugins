#!/usr/bin/env python3
"""
BCBS VT Reference Doc Styler
Applies Blue Cross VT brand styles to a pandoc reference.docx
using python-docx.

Usage: python3 style-reference-doc.py
Requires: pip install python-docx
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed.")
    print("Run: pip install python-docx")
    sys.exit(1)

# --- BCBS VT Brand Colors (2025-10 Brand Style Guide) ---
DARK_BLUE = RGBColor(0x00, 0x46, 0x7B)     # #00467B — Web Dark Blue, Title/H1 emphasis
PRIMARY_BLUE = RGBColor(0x00, 0x77, 0xC8)  # #0077C8 — Blue Cross VT Blue, PMS 3005 C (dominant)
LIGHT_BLUE = RGBColor(0x99, 0xD6, 0xEA)    # #99D6EA — Sky Blue, PMS 2975 C
DARK_GRAY = RGBColor(0x63, 0x66, 0x6A)     # #63666A — Cool Gray 10 C (H4, tertiary headings)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)         # #1a1a1a
GRAY = RGBColor(0x66, 0x66, 0x66)          # #666666

SCRIPT_DIR = Path(__file__).parent
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
OUTPUT_PATH = ASSETS_DIR / "reference.docx"

def set_paragraph_font(paragraph, name, size_pt, bold=False, color=None, italic=False):
    """Apply font settings to all runs in a paragraph, plus the paragraph format."""
    from docx.oxml.ns import qn
    # Update paragraph-level rPr via style
    pf = paragraph.paragraph_format
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color


def apply_style(doc, style_name, font_name, size_pt, bold=False, color=None,
                space_before=None, space_after=None, alignment=None):
    """Apply formatting to a named Word style."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        print(f"  Warning: Style '{style_name}' not found, skipping.")
        return

    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.bold = bold
    if color:
        font.color.rgb = color

    pf = style.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if alignment is not None:
        pf.alignment = alignment

    print(f"  ✓ Styled: {style_name} ({font_name} {size_pt}pt)")


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set page margins in inches."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
    print(f"  ✓ Margins: {top}\" top/bottom, {left}\" left, {right}\" right")


def add_footer_text(doc, text):
    """Add branded footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # Clear existing footer
        for para in footer.paragraphs:
            for run in para.runs:
                run.clear()

        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
    print(f"  ✓ Footer set")


def main():
    print("BCBS VT Reference Doc Styler")
    print("=" * 50)

    # Generate fresh pandoc reference doc if needed
    if not OUTPUT_PATH.exists():
        ASSETS_DIR.mkdir(parents=True, exist_ok=True)
        import subprocess
        print("  Generating base pandoc reference.docx...")
        result = subprocess.run(
            ["pandoc", "--print-default-data-file", "reference.docx"],
            capture_output=True
        )
        if result.returncode == 0:
            OUTPUT_PATH.write_bytes(result.stdout)
        else:
            # Create minimal docx
            doc = Document()
            doc.save(str(OUTPUT_PATH))
        print(f"  ✓ Base reference.docx created")

    print(f"\nApplying BCBS VT styles to: {OUTPUT_PATH}")
    print()

    doc = Document(str(OUTPUT_PATH))

    # --- Typography ---
    print("Applying typography...")
    # Body text — Calibri per Writing and Tone Style Guide §Text formatting
    # (Calibri / DIN 2014 / Arial are the approved correspondence fonts)
    apply_style(doc, "Normal", "Calibri", 12, color=BLACK, space_after=6)
    apply_style(doc, "Body Text", "Calibri", 12, color=BLACK, space_after=6)

    # Headings — hierarchy designed so each level is visually distinct from
    # Title and its neighbors. Size drops at each step and color pivots at
    # H2 (Dark Blue → Primary Blue) and H4 (Primary Blue → Dark Gray) keep
    # the five levels readable without violating the Brand Style Guide's
    # guidance to stay near three primary text sizes for accessibility.
    apply_style(doc, "Heading 1", "Arial", 22, bold=True, color=DARK_BLUE,
                space_before=12, space_after=6)
    apply_style(doc, "Heading 2", "Arial", 16, bold=True, color=PRIMARY_BLUE,
                space_before=10, space_after=4)
    apply_style(doc, "Heading 3", "Arial", 13, bold=True, color=PRIMARY_BLUE,
                space_before=8, space_after=4)
    apply_style(doc, "Heading 4", "Arial", 12, bold=True, color=DARK_GRAY,
                space_before=6, space_after=2)

    # Title block (pandoc renders YAML frontmatter title/author/date using
    # these styles). Title is deliberately 10pt larger than Heading 1 so
    # the document name reads as the strongest element on page one.
    apply_style(doc, "Title", "Arial", 32, bold=True, color=DARK_BLUE,
                space_before=0, space_after=12)
    apply_style(doc, "Author", "Calibri", 11, color=GRAY, space_after=2)
    apply_style(doc, "Date", "Calibri", 11, color=GRAY, space_after=12)

    # Other styles
    apply_style(doc, "Caption", "Calibri", 10, color=GRAY)
    apply_style(doc, "Intense Quote", "Calibri", 12, bold=False, color=PRIMARY_BLUE,
                space_before=8, space_after=8)

    # --- Margins ---
    print("\nSetting margins...")
    set_page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.0)

    # --- Footer ---
    print("\nAdding footer...")
    footer_text = (
        "An Independent Licensee of the Blue Cross and Blue Shield Association  |  "
        "bluecrossvt.org  |  Berlin, Vermont"
    )
    add_footer_text(doc, footer_text)

    # --- Save ---
    doc.save(str(OUTPUT_PATH))
    print(f"\n✓ Saved: {OUTPUT_PATH}")
    print()
    print("The reference.docx is ready for use with pandoc.")
    print("Run: build-letterhead.sh your-document.md output.docx")


if __name__ == "__main__":
    main()
